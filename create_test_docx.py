#!/usr/bin/env python3
"""
Create a test DOCX file with flagged language for testing.
"""

from docx import Document
from docx.shared import Inches

def create_test_docx():
    """Create a test DOCX file with various flagged terms."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Test Document for Language Flagging', 0)
    
    # Introduction paragraph with flagged terms
    intro = doc.add_paragraph()
    intro.add_run("This document contains various terms that should be flagged for inclusive language. ")
    intro.add_run("We need to address disparities in healthcare access and reduce the burden on ")
    intro.add_run("underserved communities. ")
    intro.add_run("Social determinants of health play a significant role in creating barriers ")
    intro.add_run("for vulnerable populations.")
    
    # Section 1
    doc.add_heading('Healthcare Disparities', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run("The healthcare system faces significant challenges in serving ")
    p1.add_run("marginalized groups and minority populations. ")
    p1.add_run("At-risk individuals often experience greater health-related challenges ")
    p1.add_run("due to systemic injustice and structural determinants of health.")
    
    # Section 2
    doc.add_heading('Research Methodology', level=1)
    
    p2 = doc.add_paragraph()
    p2.add_run("Our research focuses on ")
    p2.add_run("culturally relevant interventions for ")
    p2.add_run("people of color and ")
    p2.add_run("first-generation students. ")
    p2.add_run("We aim to develop ")
    p2.add_run("culturally adapted curriculum that addresses ")
    p2.add_run("social and built environments affecting ")
    p2.add_run("low-income communities.")
    
    # Section 3
    doc.add_heading('Community Engagement', level=1)
    
    p3 = doc.add_paragraph()
    p3.add_run("We will work with ")
    p3.add_run("historically marginalized communities to develop ")
    p3.add_run("inclusive research practices. ")
    p3.add_run("Our approach includes ")
    p3.add_run("advocacy for ")
    p3.add_run("social justice and addressing ")
    p3.add_run("bias in research design.")
    
    # Section 4
    doc.add_heading('Target Populations', level=1)
    
    p4 = doc.add_paragraph()
    p4.add_run("Priority populations include ")
    p4.add_run("BIPOC communities, ")
    p4.add_run("Latinx individuals, and ")
    p4.add_run("hard-to-reach populations. ")
    p4.add_run("We will focus on ")
    p4.add_run("resource-limited areas and ")
    p4.add_run("socially vulnerable communities.")
    
    # Section 5
    doc.add_heading('Implementation Strategy', level=1)
    
    p5 = doc.add_paragraph()
    p5.add_run("Our implementation will use ")
    p5.add_run("tailored interventions and ")
    p5.add_run("adapted materials for ")
    p5.add_run("underrepresented investigators. ")
    p5.add_run("We will address ")
    p5.add_run("non-medical needs and ")
    p5.add_run("built environment factors that create ")
    p5.add_run("barriers to healthcare access.")
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    p6 = doc.add_paragraph()
    p6.add_run("This comprehensive approach will help reduce ")
    p6.add_run("health disparities and improve outcomes for ")
    p6.add_run("vulnerable adults and children in ")
    p6.add_run("historically excluded communities.")
    
    # Save the document
    doc.save('test_sample.docx')
    print("Created test_sample.docx with flagged language")

if __name__ == "__main__":
    create_test_docx()
