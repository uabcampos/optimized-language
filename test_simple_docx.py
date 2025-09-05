#!/usr/bin/env python3
"""
Simple test of DOCX processing without the full pipeline
"""

import os
from docx import Document
from python_redlines.engines import XmlPowerToolsEngine

def test_simple_docx():
    """Test simple DOCX processing."""
    input_file = "test_sample.docx"
    
    print(f"Input file exists: {os.path.exists(input_file)}")
    print(f"Absolute path: {os.path.abspath(input_file)}")
    
    # Test basic DOCX reading
    try:
        doc = Document(input_file)
        print(f"✅ Successfully loaded DOCX with {len(doc.paragraphs)} paragraphs")
        
        # Test creating a modified version
        modified_doc = Document(input_file)
        modified_doc.paragraphs[0].clear()
        modified_doc.paragraphs[0].add_run("Modified text with changes")
        
        modified_path = "test_modified.docx"
        modified_doc.save(modified_path)
        print(f"✅ Created modified version: {modified_path}")
        
        # Test Python-Redlines
        engine = XmlPowerToolsEngine()
        redline_bytes, _, _ = engine.run_redline(
            author_tag="Test Author",
            original=input_file,
            modified=modified_path
        )
        
        with open("test_redlined.docx", 'wb') as f:
            f.write(redline_bytes)
        
        print("✅ Python-Redlines worked!")
        
        # Clean up
        os.remove(modified_path)
        print("✅ Cleanup completed")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    test_simple_docx()
