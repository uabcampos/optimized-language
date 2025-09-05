#!/usr/bin/env python3
"""
Simplified test using Python-Redlines for DOCX processing
"""

import os
import json
from docx import Document
from python_redlines.engines import XmlPowerToolsEngine

def create_modified_docx(input_file, suggestions):
    """Create a modified DOCX with suggestions applied."""
    doc = Document(input_file)
    
    # Apply suggestions to the document
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for original, suggestion in suggestions.items():
            if original.lower() in text.lower():
                text = text.replace(original, suggestion)
        paragraph.clear()
        paragraph.add_run(text)
    
    modified_path = "modified_" + os.path.basename(input_file)
    doc.save(modified_path)
    return modified_path

def main():
    # Load flagged terms and replacements
    with open('flagged_terms.json', 'r') as f:
        flagged_terms = json.load(f)
    
    with open('replacements.json', 'r') as f:
        replacements = json.load(f)
    
    # Create a simple mapping of terms to suggestions
    suggestions = {}
    for term in flagged_terms[:10]:  # Test with first 10 terms
        if term in replacements:
            suggestions[term] = replacements[term]
        else:
            suggestions[term] = f"[suggested: {term}]"
    
    print(f"Testing with {len(suggestions)} suggestions")
    
    # Create modified document
    input_file = "test_sample.docx"
    modified_file = create_modified_docx(input_file, suggestions)
    
    # Use Python-Redlines
    engine = XmlPowerToolsEngine()
    redline_bytes, _, _ = engine.run_redline(
        author_tag="Language Flag Tool",
        original=input_file,
        modified=modified_file
    )
    
    # Save redlined document
    output_file = "flagged_output_redlined.docx"
    with open(output_file, 'wb') as f:
        f.write(redline_bytes)
    
    print(f"✅ Created redlined document: {output_file}")
    
    # Clean up
    os.remove(modified_file)
    print("✅ Cleanup completed")

if __name__ == "__main__":
    main()
