#!/usr/bin/env python3
"""
Smart PDF language flagger with LLM-in-the-loop.

- Finds phrases from a flagged terms list and a replacements map
- Highlights or underlines in the original PDF
- Always queries the LLM for a context-aware suggestion
- If a static replacement exists, the LLM must validate or refine it
- If no static replacement exists, the LLM proposes one, keeping house rules
- Exports CSV and JSON reports with page, match, suggestion, and reasoning
- Caches LLM calls by (normalized_term, normalized_context) to reduce cost

Requires:
  pymupdf, pandas, openai, tiktoken, tenacity, python-docx
"""

import argparse
import hashlib
import json
import os
import re
import time
import multiprocessing as mp
from dataclasses import dataclass, asdict
from typing import Dict, List, Tuple, Any, Optional
from pathlib import Path
from tqdm import tqdm

import fitz  # PyMuPDF
import pandas as pd
from tenacity import retry, wait_exponential, stop_after_attempt
from dotenv import load_dotenv, find_dotenv

# DOCX support
try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.oxml.shared import qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# Pylighter support (optional)
try:
    import pylighter
    _PYLIGHTER_AVAILABLE = True
except ImportError:
    _PYLIGHTER_AVAILABLE = False

# Python-Redlines support for proper tracked changes
try:
    from python_redlines.engines import XmlPowerToolsEngine
    _REDLINES_AVAILABLE = True
except ImportError:
    _REDLINES_AVAILABLE = False

# OpenAI SDK v1.x
try:
    from openai import OpenAI
    _OPENAI_STYLE = "responses"  # prefer Responses API if available in your install
except Exception:
    OpenAI = None
    _OPENAI_STYLE = "none"

# Google Gemini API
try:
    import google.generativeai as genai
    _GEMINI_AVAILABLE = True
except ImportError:
    _GEMINI_AVAILABLE = False


# -----------------------------
# File type detection
# -----------------------------

def detect_file_type(file_path: str) -> str:
    """Detect file type based on extension."""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == '.pdf':
        return 'pdf'
    elif ext in ['.docx', '.doc']:
        if not _DOCX_AVAILABLE:
            raise RuntimeError("python-docx package not available. Install with `pip install python-docx`.")
        return 'docx'
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported types: .pdf, .docx, .doc")

# -----------------------------
# House style and prompting
# -----------------------------

SYSTEM_MSG = (
    "You are a careful, concise editor specializing in academic and scientific writing. Follow these rules:\n"
    "- Prioritize Project 2025 healthcare principles and Alabama State Bill 129 compliance.\n"
    "- Use clear, precise academic language. Prefer short, clear sentences.\n"
    "- Replace 'barrier' or 'burden' forms with 'challenge' forms.\n"
    "- Replace 'social determinants' or 'structural determinants' with "
    "  'non-medical factors such as housing, transportation, education, or nutrition'.\n"
    "- Replace standalone 'determinants' with 'influences', 'factors', or 'conditions'.\n"
    "- NEVER use the word 'determinants' in your suggestions - always replace it completely.\n"
    "- Avoid words that emphasize deficit framing. Prefer constructive, neutral alternatives.\n"
    "- Consider context carefully: preserve proper nouns, institutional names, grant titles, publication titles, and center names.\n"
    "- If the term appears to be part of a title, grant name, publication, or institutional name, suggest keeping it unchanged.\n"
    "- If a static suggestion is provided, validate or refine it for the specific context.\n"
    "- If no static suggestion exists, propose a concise replacement that preserves meaning and accuracy.\n"
    "- Keep the suggestion short, plain, and audience-centered.\n"
    "- Focus on scientific accuracy and policy alignment over style guide preferences.\n"
    "- Return strict JSON with fields: suggestion, reason."
)

USER_TEMPLATE = (
    "Term or phrase to revise: {term}\n\n"
    "Static suggestion from map (may be empty if none): {static_suggestion}\n\n"
    "Surrounding context from document:\n{context}\n\n"
    "Context Analysis: Please analyze if this term appears to be part of:\n"
    "- A proper noun, institutional name, or organization title\n"
    "- A grant title, publication title, or research project name\n"
    "- A center, department, or program name\n"
    "- A geographic location or proper name\n"
    "- A title, heading, or section header\n\n"
    "Task: Provide a replacement that fits this exact context, follows the rules, and is suitable for public-facing academic or grant text.\n"
    "If the term appears to be part of a proper noun, title, or institutional name, suggest keeping it unchanged.\n"
    "Focus on Project 2025 healthcare principles and Alabama State Bill 129 compliance.\n"
    "Return JSON only with fields: suggestion, reason."
)

def load_env(env_path: Optional[str] = None) -> None:
    """
    Load environment variables from a .env file if present.
    If env_path is provided, load from that path; otherwise search upwards.
    """
    if env_path:
        load_dotenv(dotenv_path=env_path, override=False)
    else:
        # find_dotenv searches up the directory tree; only load if found
        dotenv_file = find_dotenv(usecwd=True)
        if dotenv_file:
            load_dotenv(dotenv_file, override=False)

# -----------------------------
# Data classes
# -----------------------------

def should_skip_term(term: str, skip_terms: List[str]) -> bool:
    """Check if a term should be skipped based on skip terms list and their variations."""
    if not skip_terms:
        return False
    
    term_lower = term.lower()
    
    for skip_term in skip_terms:
        skip_lower = skip_term.lower()
        
        # Exact match
        if term_lower == skip_lower:
            return True
        
        # Check for word variations (more precise matching)
        # Only match if the term is a variation of the skip term, not just starts with it
        common_suffixes = ['s', 'es', 'ed', 'ing', 'ly', 'tion', 'sion', 'ness', 'ment', 'able', 'ible', 'ful', 'less']
        
        # Check if term is skip_term + suffix
        for suffix in common_suffixes:
            if term_lower == skip_lower + suffix:
                return True
        
        # Check if skip_term is term + suffix
        for suffix in common_suffixes:
            if skip_lower == term_lower + suffix:
                return True
        
        # Check if they share the same root (both have suffixes)
        for suffix in common_suffixes:
            if term_lower.endswith(suffix) and skip_lower.endswith(suffix):
                term_root = term_lower[:-len(suffix)]
                skip_root = skip_lower[:-len(suffix)]
                if term_root == skip_root:
                    return True
    
    return False

def is_proper_noun_context(text: str, match_start: int, match_end: int) -> bool:
    """Check if the matched term is part of a proper noun (like department names, titles, grants, publications)."""
    # Get context around the match
    context_start = max(0, match_start - 100)
    context_end = min(len(text), match_end + 100)
    context = text[context_start:context_end]
    
    # Check for common proper noun patterns
    proper_noun_patterns = [
        # Institutional names
        r'\b(Department|Center|Institute|School|College|University|Hospital|Clinic|Program|Office|Division|Unit|Section|Group|Team|Committee|Board|Council|Foundation|Organization|Association|Society|Academy|Laboratory|Lab)\s+of\s+',
        r'\b(Department|Center|Institute|School|College|University|Hospital|Clinic|Program|Office|Division|Unit|Section|Group|Team|Committee|Board|Council|Foundation|Organization|Association|Society|Academy|Laboratory|Lab)\s+for\s+',
        r'\b(Department|Center|Institute|School|College|University|Hospital|Clinic|Program|Office|Division|Unit|Section|Group|Team|Committee|Board|Council|Foundation|Organization|Association|Society|Academy|Laboratory|Lab)\s+',
        # Titles and positions
        r'\b(Chair|Director|Dean|President|Vice|Executive|Senior|Chief|Head|Lead|Principal|Coordinator|Manager|Administrator)\s+of\s+',
        r'\b(Chair|Director|Dean|President|Vice|Executive|Senior|Chief|Head|Lead|Principal|Coordinator|Manager|Administrator)\s+for\s+',
        r'\b(Chair|Director|Dean|President|Vice|Executive|Senior|Chief|Head|Lead|Principal|Coordinator|Manager|Administrator)\s+',
        # Grant and funding patterns
        r'\b(Grant|Award|Contract|Cooperative Agreement|R01|R21|R03|P01|P50|U01|U19|K01|K08|K23|F31|F32|T32|T35)\s+',
        r'\b(NIH|NSF|CDC|AHRQ|PCORI|NCI|NHLBI|NIDDK|NIMH|NIA|NIAID|NINDS|NICHD|NIEHS|NIGMS|NIDCR|NIDCD|NIDCR|NLM|NCCIH|NIMHD|NINR|NIBIB|NIDA|NIAAA|NICHD|NIEHS|NIGMS|NIDCR|NIDCD|NIDCR|NLM|NCCIH|NIMHD|NINR|NIBIB|NIDA|NIAAA)\s+',
        # Publication patterns
        r'\b(Journal|Journal of|American Journal|British Journal|New England Journal|Lancet|Nature|Science|Cell|PNAS|PLOS|BMC|Frontiers)\s+',
        r'\b(Volume|Vol\.|Issue|No\.|Pages|pp\.|DOI|PMID|PMCID)\s+',
        # Title patterns (often in quotes or italics)
        r'["\']([^"\']*' + re.escape(text[match_start:match_end]) + r'[^"\']*)["\']',
        r'\b(Title|Abstract|Introduction|Methods|Results|Discussion|Conclusion|Background|Objective|Purpose|Aim)\s*:',
        # Center and program names
        r'\b(Center for|Program for|Initiative for|Project for|Study of|Research on|Institute for|Laboratory for)\s+',
        # Geographic and proper names
        r'\b(Alabama|Birmingham|Montgomery|Huntsville|Mobile|Tuscaloosa|Auburn|Troy|Jacksonville|Florence|Gadsden|Dothan|Decatur|Phenix City|Prattville|Opelika|Anniston|Bessemer|Hoover|Vestavia Hills|Homewood|Mountain Brook|Trussville|Pelham|Helena|Calera|Chelsea|Leeds|Moody|Pell City|Sylacauga|Alexander City|Talladega|Oxford|Albertville|Fort Payne|Scottsboro|Athens|Hartselle|Cullman|Jasper|Hamilton|Winfield|Fayette|Carrollton|Livingston|Demopolis|Selma|Marion|Greensboro|Eutaw|Tuskegee|Union Springs|Auburn|Opelika|Phenix City|Columbus|Valley|Lanett|LaGrange|West Point|Auburn|Opelika|Phenix City|Columbus|Valley|Lanett|LaGrange|West Point)\s+',
    ]
    
    for pattern in proper_noun_patterns:
        if re.search(pattern, context, re.IGNORECASE):
            return True
    
    # Check for title-like patterns (capitalized words, quotes, etc.)
    # Look for patterns that suggest this is part of a title
    title_indicators = [
        r'^[A-Z][^.!?]*$',  # Entire line is capitalized (title case)
        r'["\'].*["\']',    # Text in quotes
        r'\b[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+',  # Multiple capitalized words
        r'\b(Figure|Table|Appendix|Chapter|Section)\s+\d+',  # Figure/table references
    ]
    
    for pattern in title_indicators:
        if re.search(pattern, context, re.IGNORECASE):
            return True
    
    return False

@dataclass
class Hit:
    page_num: int
    original_key: str          # flagged term as defined in list or map
    matched_text: str          # actual matched text on page
    suggestion: str
    reason: str
    bbox: Tuple[float, float, float, float]  # union of matched words
    context: str

@dataclass
class DocumentAnalysis:
    """Comprehensive analysis of a document's themes, objectives, and strategic alignment."""
    document_type: str  # "research", "clinical", "community_health", "policy", "other"
    main_themes: List[str]  # Key themes and topics
    scientific_focus: List[str]  # Scientific concepts and methodologies
    target_population: str  # Primary target population
    intervention_approach: str  # Type of intervention or approach
    key_objectives: List[str]  # Main objectives and goals
    methodology: str  # Research or implementation methodology
    policy_relevance: List[str]  # Policy-relevant themes
    nih_priority_alignment: List[str]  # Alignment with current NIH Director priorities
    alignment_opportunities: List[str]  # Areas for Project 2025/Bill 129 alignment
    strategic_recommendations: List[str]  # High-level recommendations
    nih_alignment_score: float  # 0-100% alignment with NIH priorities
    project_2025_score: float  # 0-100% alignment with Project 2025
    alabama_sb129_score: float  # 0-100% compliance with Alabama SB 129

@dataclass
class StrategicRecommendation:
    """Individual strategic recommendation for document improvement."""
    category: str  # "reframing", "language", "content", "structure"
    priority: str  # "high", "medium", "low"
    title: str  # Brief title of the recommendation
    description: str  # Detailed description
    rationale: str  # Why this recommendation matters
    implementation: str  # How to implement the recommendation
    expected_impact: str  # Expected impact on alignment/compliance

@dataclass
class AlignmentSuggestion:
    """Specific suggestion for aligning with Project 2025 or Alabama Bill 129."""
    policy_source: str  # "Project 2025" or "Alabama Bill 129"
    principle: str  # Specific principle or requirement
    current_state: str  # How the document currently addresses this
    suggested_improvement: str  # Specific improvement suggestion
    implementation_guidance: str  # How to implement the improvement
    priority: str  # "critical", "important", "nice_to_have"

# -----------------------------
# Helpers
# -----------------------------

def load_json_file(path: str) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def normalize_token(t: str) -> str:
    # keep letters, digits, internal hyphens and apostrophes
    return re.sub(r"[^\w\-’']", "", t).lower()

def words_by_order(page: fitz.Page):
    words = page.get_text("words")  # (x0, y0, x1, y1, text, block, line, word)
    words.sort(key=lambda w: (w[5], w[6], w[7], w[1], w[0]))
    return words

def get_union_bbox(words):
    x0 = min(w[0] for w in words)
    y0 = min(w[1] for w in words)
    x1 = max(w[2] for w in words)
    y1 = max(w[3] for w in words)
    return (x0, y0, x1, y1)

def join_words(words):
    return " ".join(w[4] for w in words)

def build_match_list(flagged_terms: List[str], repl_map: Dict[str, str]) -> List[str]:
    # union, dedup, sort by length desc so longer phrases match first
    all_terms = set(t.strip() for t in flagged_terms) | set(k.strip() for k in repl_map.keys())
    all_terms = [t for t in all_terms if t]
    all_terms.sort(key=lambda s: len(s), reverse=True)
    return all_terms

def build_phrase_tokens(terms: List[str]) -> List[Tuple[str, List[str]]]:
    out = []
    for term in terms:
        toks = [normalize_token(t) for t in term.split()]
        toks = [t for t in toks if t]
        if toks:
            out.append((term, toks))
    # already sorted by length from caller
    return out

def extract_context_tokens(orig_tokens: List[str], start: int, end: int, window: int = 12) -> str:
    left = max(0, start - window)
    right = min(len(orig_tokens), end + window)
    return " ".join(orig_tokens[left:right])

def term_context_cache_key(term: str, context: str) -> str:
    key = f"{term.lower()}||{context.strip().lower()}"
    return hashlib.sha256(key.encode("utf-8")).hexdigest()

# -----------------------------
# OpenAI client and call
# -----------------------------

def get_openai_client() -> OpenAI:
    if OpenAI is None:
        raise RuntimeError("openai package not available. Install with `pip install openai`.")
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set. Set it in your shell or supply a .env file (see --env-file).")
    return OpenAI(api_key=api_key)

def get_gemini_client(model: str = "gemini-1.5-flash"):
    """Initialize Gemini client."""
    if not _GEMINI_AVAILABLE:
        raise RuntimeError("google-generativeai package not available. Install with `pip install google-generativeai`.")
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set. Set it in your shell or supply a .env file (see --env-file).")
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(model)

@retry(wait=wait_exponential(multiplier=1, min=1, max=20), stop=stop_after_attempt(4))
def llm_suggest_gemini(model, term: str, static_suggestion: str, context: str, temperature: float) -> Tuple[str, str]:
    """
    Returns (suggestion, reason) JSON-parsed from Gemini model output.
    """
    user_msg = USER_TEMPLATE.format(term=term, static_suggestion=static_suggestion or "", context=context)
    
    # Enhanced prompt for Gemini to ensure JSON output
    enhanced_system_msg = f"""{SYSTEM_MSG}

IMPORTANT: You must respond with valid JSON only. Use this exact format:
{{
  "suggestion": "your suggestion here",
  "reason": "your reason here"
}}

Do not include any other text, explanations, or formatting outside the JSON."""
    
    full_prompt = f"{enhanced_system_msg}\n\n{user_msg}"
    
    try:
        response = model.generate_content(
            full_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                max_output_tokens=500,
            )
        )
        content = response.text.strip()
        
        # Clean up the response to extract JSON
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()
        
        # Try to find JSON in the response
        if "{" in content and "}" in content:
            start = content.find("{")
            end = content.rfind("}") + 1
            content = content[start:end]
            
    except Exception as e:
        raise RuntimeError(f"Gemini API error: {e}")

    try:
        data = json.loads(content)
        suggestion = data.get("suggestion", "").strip()
        reason = data.get("reason", "").strip()
        if not suggestion:
            raise ValueError("Empty suggestion from model.")
        return suggestion, reason
    except Exception as e:
        # Safe fallback: if parsing fails, use static suggestion or echo term
        fallback = static_suggestion or term
        return fallback, f"Used fallback suggestion due to parsing issue: {str(e)[:100]}"

@retry(wait=wait_exponential(multiplier=1, min=1, max=20), stop=stop_after_attempt(4))
def llm_suggest_openai(client: OpenAI, model: str, term: str, static_suggestion: str, context: str, temperature: float) -> Tuple[str, str]:
    """
    Returns (suggestion, reason) JSON-parsed from OpenAI model output.
    """
    user_msg = USER_TEMPLATE.format(term=term, static_suggestion=static_suggestion or "", context=context)

    # Prefer the Responses API if your SDK supports it
    try:
        resp = client.responses.create(
            model=model,
            temperature=temperature,
            input=[
                {"role": "system", "content": SYSTEM_MSG},
                {"role": "user", "content": user_msg},
            ],
            response_format={"type": "json_object"},
        )
        content = resp.output_text
    except Exception:
        # Fallback to chat.completions if Responses not present
        chat = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[
                {"role": "system", "content": SYSTEM_MSG},
                {"role": "user", "content": user_msg},
            ],
        )
        content = chat.choices[0].message.content

    try:
        data = json.loads(content)
        suggestion = data.get("suggestion", "").strip()
        reason = data.get("reason", "").strip()
        if not suggestion:
            raise ValueError("Empty suggestion from model.")
        return suggestion, reason
    except Exception as e:
        # Safe fallback: if parsing fails, use static suggestion or echo term
        fallback = static_suggestion or term
        return fallback, "Used fallback suggestion due to parsing issue."

def llm_suggest(client, model: str, term: str, static_suggestion: str, context: str, temperature: float, api_type: str = "auto") -> Tuple[str, str]:
    """
    Universal LLM suggest function that works with both OpenAI and Gemini.
    """
    if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
        return llm_suggest_gemini(client, term, static_suggestion, context, temperature)
    else:
        return llm_suggest_openai(client, model, term, static_suggestion, context, temperature)

def analyze_document_themes(pdf_text: str = None, docx_text: str = None, 
                          model: str = "gemini-1.5-flash", 
                          api_type: str = "auto") -> DocumentAnalysis:
    """
    Analyze document themes, objectives, and strategic alignment opportunities.
    
    Args:
        pdf_text: Extracted text from PDF
        docx_text: Extracted text from DOCX
        model: LLM model to use for analysis
        api_type: API provider type
        
    Returns:
        DocumentAnalysis object with comprehensive document insights
    """
    # Combine text sources
    full_text = ""
    if pdf_text:
        full_text += pdf_text
    if docx_text:
        full_text += docx_text
    
    if not full_text.strip():
        return DocumentAnalysis(
            document_type="unknown",
            main_themes=[],
            scientific_focus=[],
            target_population="unknown",
            intervention_approach="unknown",
            key_objectives=[],
            methodology="unknown",
            policy_relevance=[],
            nih_priority_alignment=[],
            alignment_opportunities=[],
            strategic_recommendations=[],
            nih_alignment_score=0.0,
            project_2025_score=0.0,
            alabama_sb129_score=0.0
        )
    
    # Truncate text if too long (keep first 8000 characters for analysis)
    analysis_text = full_text[:8000] if len(full_text) > 8000 else full_text
    
    # Create LLM client
    try:
        if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
            client = get_gemini_client(model)
        else:
            client = get_openai_client()
    except Exception as e:
        print(f"Error creating client for document analysis: {e}")
        return DocumentAnalysis(
            document_type="unknown",
            main_themes=[],
            scientific_focus=[],
            target_population="unknown",
            intervention_approach="unknown",
            key_objectives=[],
            methodology="unknown",
            policy_relevance=[],
            nih_priority_alignment=[],
            alignment_opportunities=[],
            strategic_recommendations=[],
            nih_alignment_score=0.0,
            project_2025_score=0.0,
            alabama_sb129_score=0.0
        )
    
    # Generate analysis using LLM
    analysis_prompt = f"""
    Analyze the following document and provide a comprehensive analysis in JSON format. Focus on understanding the document's purpose, themes, and potential alignment with current NIH priorities and healthcare policy.

    Document Text:
    {analysis_text}

    Please provide your analysis in this exact JSON format:
    {{
        "document_type": "research|clinical|community_health|policy|other",
        "main_themes": ["theme1", "theme2", "theme3"],
        "scientific_focus": ["concept1", "concept2", "concept3"],
        "target_population": "description of primary target population",
        "intervention_approach": "description of intervention or approach",
        "key_objectives": ["objective1", "objective2", "objective3"],
        "methodology": "description of research or implementation methodology",
        "policy_relevance": ["policy_theme1", "policy_theme2"],
        "nih_priority_alignment": ["priority1", "priority2", "priority3"],
        "alignment_opportunities": ["opportunity1", "opportunity2"],
        "strategic_recommendations": ["recommendation1", "recommendation2"],
        "nih_alignment_score": 75.5,
        "project_2025_score": 68.2,
        "alabama_sb129_score": 72.1
    }}

    Current NIH Director Priorities (August 2024) - Evaluate alignment with:
    1. **Chronic Disease and Nutrition**: Research on poor diets causing chronic conditions, healthy diet prevention/management, childhood obesity, insulin resistance
    2. **Replication and Reproducibility**: Robust, reproducible results, replication studies, negative findings publication
    3. **Real-World Data Platform**: Secure national infrastructure, data integration, privacy-focused computational analysis
    4. **Artificial Intelligence (AI)**: AI Strategic Plan, transparency, replication standards, patient benefit translation
    5. **Autism**: Etiology, treatment, care needs, data science initiatives, data gap identification
    6. **Alternative Testing Models**: Move away from animal testing, pharmaceutical alternative models
    7. **Health Outcomes and Disparities**: Scientifically valid, measurable outcomes, solution-oriented approaches, measurable factors (redlining, housing discrimination)
    8. **Training Future Biomedical Scientists**: Merit-based programs, American biomedical research leadership
    9. **HIV/AIDS Research**: Implementation science, reaching at-risk communities
    10. **Foreign Research Oversight**: Responsible use of taxpayer dollars, improved oversight

    Focus on:
    1. Identifying the document type and main purpose
    2. Extracting key scientific themes and concepts
    3. Understanding the target population and intervention approach
    4. Identifying policy-relevant themes
    5. Evaluating alignment with specific NIH priorities listed above
    6. Suggesting alignment opportunities with NIH priorities
    7. Providing strategic recommendations for improvement
    8. Scoring alignment (0-100) with NIH priorities
    9. Scoring alignment (0-100) with Project 2025 healthcare principles
    10. Scoring compliance (0-100) with Alabama State Bill 129

    Respond with valid JSON only.
    """
    
    try:
        if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
            response = client.generate_content(
                analysis_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=2000,
                )
            )
            content = response.text.strip()
        else:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an expert document analyst specializing in healthcare policy and grant writing. Provide detailed, accurate analysis in JSON format."},
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            content = response.choices[0].message.content.strip()
        
        # Parse JSON response
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()
        
        # Extract JSON
        if "{" in content and "}" in content:
            start = content.find("{")
            end = content.rfind("}") + 1
            content = content[start:end]
        
        data = json.loads(content)
        
        return DocumentAnalysis(
            document_type=data.get("document_type", "unknown"),
            main_themes=data.get("main_themes", []),
            scientific_focus=data.get("scientific_focus", []),
            target_population=data.get("target_population", "unknown"),
            intervention_approach=data.get("intervention_approach", "unknown"),
            key_objectives=data.get("key_objectives", []),
            methodology=data.get("methodology", "unknown"),
            policy_relevance=data.get("policy_relevance", []),
            nih_priority_alignment=data.get("nih_priority_alignment", []),
            alignment_opportunities=data.get("alignment_opportunities", []),
            strategic_recommendations=data.get("strategic_recommendations", []),
            nih_alignment_score=float(data.get("nih_alignment_score", 0.0)),
            project_2025_score=float(data.get("project_2025_score", 0.0)),
            alabama_sb129_score=float(data.get("alabama_sb129_score", 0.0))
        )
        
    except Exception as e:
        print(f"Error in document analysis: {e}")
        return DocumentAnalysis(
            document_type="unknown",
            main_themes=[],
            scientific_focus=[],
            target_population="unknown",
            intervention_approach="unknown",
            key_objectives=[],
            methodology="unknown",
            policy_relevance=[],
            nih_priority_alignment=[],
            alignment_opportunities=[],
            strategic_recommendations=[],
            nih_alignment_score=0.0,
            project_2025_score=0.0,
            alabama_sb129_score=0.0
        )

# -----------------------------
# Core matching and annotation
# -----------------------------

def process_terms_chunk(args) -> List[Hit]:
    """Worker function for processing a chunk of terms with proper text matching and bbox calculation."""
    page_text, phrase_tokens, repl_map, model, temperature, page_num, page_words, api_type, skip_terms = args
    
    # Create a new client for this process
    try:
        if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
            client = get_gemini_client(model)
        else:
            client = get_openai_client()
    except Exception as e:
        print(f"Error creating client in worker process: {e}")
        return []
    
    cache = {}  # Each process gets its own cache
    
    hits = []
    skipped_count = 0
    processed_count = 0
    norm_tokens = [normalize_token(w[4]) for w in page_words]
    orig_tokens = [w[4] for w in page_words]
    
    print(f"🔍 Processing {len(phrase_tokens)} terms on page {page_num} with skip terms: {skip_terms}")
    
    # Keep track of already consumed word indices to avoid overlaps
    used: List[int] = []
    
    def span_free(span: List[int]) -> bool:
        return all(idx not in used for idx in span)
    
    for phrase, toks in phrase_tokens:
        if not toks:
            continue
            
        n = len(toks)
        i = 0
        while i <= len(norm_tokens) - n:
            window = norm_tokens[i:i+n]
            if window == toks and span_free(list(range(i, i+n))):
                span = list(range(i, i+n))
                matched_words = [page_words[k] for k in span]
                matched_text = join_words(matched_words)
                bbox = get_union_bbox(matched_words)
                context = extract_context_tokens(orig_tokens, i, i+n)
                
                # Check if this is part of a proper noun context
                if is_proper_noun_context(context, 0, len(matched_text)):
                    i += 1
                    continue
                
                # Check if this term should be skipped
                if should_skip_term(phrase, skip_terms):
                    print(f"⏭️  Skipping term: '{phrase}' (matches skip terms: {skip_terms})")
                    skipped_count += 1
                    i += 1
                    continue
                
                # Get static suggestion if available
                static_suggestion = repl_map.get(phrase, "")
                
                # Always ask LLM
                cache_key = term_context_cache_key(phrase, context)
                if cache_key in cache:
                    suggestion, reason = cache[cache_key]
                else:
                    suggestion, reason = llm_suggest(client, model, phrase, static_suggestion, context, temperature, api_type)
                    cache[cache_key] = (suggestion, reason)
                
                hit = Hit(
                    page_num=page_num,
                    original_key=phrase,
                    matched_text=matched_text,
                    suggestion=suggestion,
                    reason=reason,
                    bbox=bbox,  # Proper bbox calculation
                    context=context
                )
                hits.append(hit)
                processed_count += 1
                used.extend(span)
                i += n
                continue
            i += 1
    
    print(f"📊 Page {page_num} summary: {processed_count} processed, {skipped_count} skipped, {len(hits)} hits found")
    return hits

def find_hits_on_page(page: fitz.Page,
                      phrase_tokens: List[Tuple[str, List[str]]],
                      repl_map: Dict[str, str],
                      client,
                      model: str,
                      temperature: float,
                      cache: Dict[str, Tuple[str, str]],
                      api_type: str = "auto") -> List[Hit]:
    hits: List[Hit] = []
    words = words_by_order(page)
    norm_tokens = [normalize_token(w[4]) for w in words]
    orig_tokens = [w[4] for w in words]

    # Keep track of already consumed word indices to avoid overlaps
    used: List[int] = []

    def span_free(span: List[int]) -> bool:
        return all(idx not in used for idx in span)

    for phrase_idx, (phrase, toks) in enumerate(phrase_tokens):
        n = len(toks)
        if n == 0:
            continue
        i = 0
        while i <= len(norm_tokens) - n:
            window = norm_tokens[i:i+n]
            if window == toks and span_free(list(range(i, i+n))):
                span = list(range(i, i+n))
                matched_words = [words[k] for k in span]
                matched_text = join_words(matched_words)
                bbox = get_union_bbox(matched_words)
                context = extract_context_tokens(orig_tokens, i, i+n)

                # Check if this is part of a proper noun context
                if is_proper_noun_context(context, 0, len(matched_text)):
                    continue  # Skip this match if it's part of a proper noun
                
                static_suggestion = repl_map.get(phrase, "")
                # Always ask LLM
                cache_key = term_context_cache_key(phrase, context)
                if cache_key in cache:
                    suggestion, reason = cache[cache_key]
                else:
                    suggestion, reason = llm_suggest(client, model, phrase, static_suggestion, context, temperature, api_type)
                    cache[cache_key] = (suggestion, reason)

                hit = Hit(
                    page_num=page.number + 1,
                    original_key=phrase,
                    matched_text=matched_text,
                    suggestion=suggestion,
                    reason=reason,
                    bbox=bbox,
                    context=context
                )
                hits.append(hit)
                used.extend(span)
                i += n
                continue
            i += 1

    return hits

def annotate_hit(page: fitz.Page, hit: Hit, style: str = "highlight"):
    # Build quads for the match to handle line wraps
    hb = fitz.Rect(*hit.bbox)
    page_words = words_by_order(page)
    quads = []
    for w in page_words:
        # Validate rectangle coordinates before creating rect
        x0, y0, x1, y1 = w[0], w[1], w[2], w[3]
        if x0 < x1 and y0 < y1 and x0 >= 0 and y0 >= 0:  # Valid rectangle
            try:
                rect = fitz.Rect(x0, y0, x1, y1)
                if rect.intersects(hb):
                    quad = fitz.Quad(rect)
                    quads.append(quad)
            except (ValueError, TypeError) as e:
                # Skip invalid rectangles
                continue

    # Fall back to bbox if no quads are picked up
    if not quads:
        target = hb
    else:
        target = quads

    if style == "underline":
        annot = page.add_underline_annot(target)
    else:
        annot = page.add_highlight_annot(target)

    # Set tooltip information
    annot.set_info({
        "title": "Flagged Language",
        "subject": "Suggested revision",
        "content": f"Original: {hit.matched_text}\nSuggest: {hit.suggestion}\nWhy: {hit.reason[:350]}"
    })
    
    # Add a text annotation for better tooltip visibility
    # Position the text annotation slightly above the highlight
    text_rect = fitz.Rect(hb.x0, hb.y0 - 15, hb.x0 + 200, hb.y0 - 5)
    text_annot = page.add_text_annot(text_rect.tl, f"💡 {hit.suggestion}")
    text_annot.set_info({
        "title": "Language Suggestion",
        "subject": "Rewrite Suggestion", 
        "content": f"Original: {hit.matched_text}\n\nSuggested: {hit.suggestion}\n\nReason: {hit.reason[:200]}..."
    })
    
    # Set colors to make annotations more visible
    annot.set_colors(stroke=fitz.utils.getColor("red"))
    annot.set_opacity(0.3)  # Semi-transparent highlight
    text_annot.set_colors(stroke=fitz.utils.getColor("blue"))
    
    annot.update()
    text_annot.update()

# -----------------------------
# DOCX Processing Functions
# -----------------------------

def extract_docx_text_with_positions(doc: Document) -> List[Dict]:
    """Extract text from DOCX with detailed position information.
    Returns list of text segments with their positions."""
    text_segments = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        if para_text.strip():
            text_segments.append({
                'text': para_text,
                'para_idx': para_idx,
                'paragraph': paragraph
            })
    
    return text_segments

def find_docx_matches_advanced(text_segments: List[Dict], flagged_terms: List[str]) -> List[Dict]:
    """Find matches in DOCX text with better context extraction."""
    matches = []
    
    for segment in text_segments:
        text = segment['text']
        text_lower = text.lower()
        para_idx = segment['para_idx']
        
        for term in flagged_terms:
            term_lower = term.lower()
            if term_lower in text_lower:
                # Find all occurrences of the term
                start = 0
                while True:
                    pos = text_lower.find(term_lower, start)
                    if pos == -1:
                        break
                    
                    # Get context around the match
                    context_start = max(0, pos - 100)
                    context_end = min(len(text), pos + len(term) + 100)
                    context = text[context_start:context_end]
                    
                    matches.append({
                        'term': term,
                        'matched_text': text[pos:pos + len(term)],
                        'para_idx': para_idx,
                        'start_pos': pos,
                        'end_pos': pos + len(term),
                        'context': context,
                        'paragraph': segment['paragraph']
                    })
                    
                    start = pos + 1
    
    return matches

def create_track_changes_docx(original_doc: Document, hits: List['Hit']) -> Document:
    """Create a new document with track changes showing suggested edits."""
    # Create a copy of the original document
    new_doc = Document()
    
    # Copy all paragraphs with modifications
    for para_idx, paragraph in enumerate(original_doc.paragraphs):
        new_para = new_doc.add_paragraph()
        
        # Get all hits for this paragraph
        para_hits = [hit for hit in hits if hit.page_num - 1 == para_idx]
        
        if para_hits:
            # Process paragraph with track changes
            para_text = paragraph.text
            modifications = []
            
            for hit in para_hits:
                # Find the position of the matched text
                pos = para_text.lower().find(hit.matched_text.lower())
                if pos != -1:
                    modifications.append({
                        'start': pos,
                        'end': pos + len(hit.matched_text),
                        'original': hit.matched_text,
                        'suggestion': hit.suggestion,
                        'reason': hit.reason
                    })
            
            # Sort modifications by position (reverse order to avoid index issues)
            modifications.sort(key=lambda x: x['start'], reverse=True)
            
            # Apply modifications
            current_text = para_text
            for mod in modifications:
                before = current_text[:mod['start']]
                after = current_text[mod['end']:]
                current_text = before + f"[{mod['original']} → {mod['suggestion']}]" + after
            
            # Add the modified text
            new_para.add_run(current_text)
            
            # Add suggestions as inline text (since comments aren't available)
            if modifications:
                suggestion_text = " | ".join([f"{mod['original']}→{mod['suggestion']}" for mod in modifications])
                new_para.add_run(f" [SUGGESTIONS: {suggestion_text}]")
        else:
            # Copy paragraph as-is
            new_para.add_run(paragraph.text)
    
    return new_doc

def highlight_docx_simple(input_docx: str, hits: List['Hit'], output_docx: str) -> None:
    """Simple highlighting approach for DOCX files."""
    try:
        # Load the document
        doc = Document(input_docx)
        
        # Create a mapping of terms to their suggestions
        term_suggestions = {}
        for hit in hits:
            if hit.original_key not in term_suggestions:
                term_suggestions[hit.original_key] = hit.suggestion
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Find and highlight flagged terms
                for term, suggestion in term_suggestions.items():
                    if term.lower() in paragraph.text.lower():
                        # Split the paragraph into runs and highlight
                        runs = paragraph.runs
                        if runs:
                            # Get the full text
                            full_text = paragraph.text
                            
                            # Find the term position
                            pos = full_text.lower().find(term.lower())
                            if pos != -1:
                                # Clear existing runs
                                paragraph.clear()
                                
                                # Add text before the term
                                if pos > 0:
                                    paragraph.add_run(full_text[:pos])
                                
                                # Add highlighted term
                                highlighted_run = paragraph.add_run(full_text[pos:pos + len(term)])
                                highlighted_run.bold = True
                                # Use a valid highlight color index
                                from docx.enum.text import WD_COLOR_INDEX
                                highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                
                                # Add text after the term
                                if pos + len(term) < len(full_text):
                                    paragraph.add_run(full_text[pos + len(term):])
                                
                                # Add suggestion as inline text
                                suggestion_run = paragraph.add_run(f" [→{suggestion}]")
                                suggestion_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                                suggestion_run.italic = True
        
        # Save the highlighted document
        doc.save(output_docx)
        
    except Exception as e:
        print(f"Warning: Simple highlighting failed: {e}")
        # Fallback: copy original file
        import shutil
        shutil.copy2(input_docx, output_docx)

def process_docx_advanced(input_docx: str, flagged_terms: List[str], repl_map: Dict[str, str], 
                         model: str, temperature: float, cache: Dict[str, Tuple[str, str]], api_type: str = "auto") -> Tuple[Document, List['Hit']]:
    """Advanced DOCX processing with track changes and highlighting."""
    # Use absolute path to avoid issues
    input_path = os.path.abspath(input_docx)
    doc = Document(input_path)
    hits = []
    
    # Extract text data
    text_segments = extract_docx_text_with_positions(doc)
    
    # Find matches
    matches = find_docx_matches_advanced(text_segments, flagged_terms)
    
    # Process each match
    for match in matches:
        term = match['term']
        matched_text = match['matched_text']
        context = match['context']
        para_idx = match['para_idx']
        
        # Get static suggestion if available
        static_suggestion = repl_map.get(term, "")
        
        # Get LLM suggestion
        cache_key = f"{term.lower()}|{context.lower()[:100]}"
        if cache_key in cache:
            suggestion, reason = cache[cache_key]
        else:
            if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
                client = get_gemini_client(model)
            else:
                client = get_openai_client()
            suggestion, reason = llm_suggest(client, model, term, static_suggestion, context, temperature, api_type)
            cache[cache_key] = (suggestion, reason)
        
        # Create hit object
        hit = Hit(
            page_num=para_idx + 1,  # Use paragraph number as "page"
            original_key=term,
            matched_text=matched_text,
            suggestion=suggestion,
            reason=reason,
            bbox=(0, 0, 0, 0),  # No bbox for DOCX
            context=context
        )
        hits.append(hit)
    
    return doc, hits

def create_modified_docx_with_suggestions(input_docx: str, hits: List['Hit']) -> str:
    """Create a modified DOCX with suggested changes applied."""
    # Use absolute path to avoid issues
    input_path = os.path.abspath(input_docx)
    doc = Document(input_path)
    
    # Group hits by paragraph
    para_hits = {}
    for hit in hits:
        para_idx = hit.page_num - 1  # Convert to 0-based index
        if para_idx not in para_hits:
            para_hits[para_idx] = []
        para_hits[para_idx].append(hit)
    
    # Process each paragraph with hits
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if para_idx in para_hits:
            hits_for_para = para_hits[para_idx]
            
            # Get the original text
            original_text = paragraph.text
            
            # Apply suggestions (replace flagged terms with suggestions)
            modified_text = original_text
            for hit in hits_for_para:
                # Replace the flagged term with the suggestion
                modified_text = modified_text.replace(hit.matched_text, hit.suggestion)
            
            # Clear the paragraph and add the modified text
            paragraph.clear()
            paragraph.add_run(modified_text)
    
    # Save the modified document to a safe location
    import tempfile
    modified_path = os.path.join(tempfile.gettempdir(), f"modified_{os.path.basename(input_docx)}")
    doc.save(modified_path)
    return modified_path

def process_docx_with_redlines(input_docx: str, hits: List['Hit'], outdir: str) -> str:
    """Use Python-Redlines to create proper tracked changes document."""
    if not _REDLINES_AVAILABLE:
        print("Warning: Python-Redlines not available, falling back to simple highlighting")
        return None
    
    try:
        # Use absolute paths and ensure they exist
        input_path = os.path.abspath(input_docx)
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Create modified version with suggestions
        modified_docx = create_modified_docx_with_suggestions(input_path, hits)
        if not os.path.exists(modified_docx):
            raise FileNotFoundError(f"Modified file not created: {modified_docx}")
        
        # Use Python-Redlines to create tracked changes
        engine = XmlPowerToolsEngine()
        redlined_docx = os.path.join(outdir, "flagged_output_redlined.docx")
        
        print(f"Creating tracked changes: {input_path} -> {modified_docx}")
        
        # Run the redline comparison
        redline_bytes, _, _ = engine.run_redline(
            author_tag="Language Flag Tool",
            original=input_path,
            modified=modified_docx
        )
        
        # Save the redline bytes to file
        with open(redlined_docx, 'wb') as f:
            f.write(redline_bytes)
        
        print(f"Successfully created tracked changes: {redlined_docx}")
        
        # Clean up the temporary modified file
        if os.path.exists(modified_docx):
            os.remove(modified_docx)
        
        return redlined_docx
        
    except Exception as e:
        print(f"Warning: Python-Redlines failed: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_docx_file(input_docx: str,
                     flagged_terms: List[str],
                     repl_map: Dict[str, str],
                     outdir: str,
                     model: str,
                     temperature: float,
                     api_type: str = "auto",
                     skip_terms: List[str] = None) -> Tuple[str, List[Hit]]:
    """Process DOCX file with Python-Redlines tracked changes and highlighting."""
    os.makedirs(outdir, exist_ok=True)
    
    # Build search index
    all_terms = build_match_list(flagged_terms, repl_map)
    
    # Process document with advanced method
    cache = {}
    doc, hits = process_docx_advanced(input_docx, all_terms, repl_map, model, temperature, cache, api_type)
    
    # Create output versions
    out_docx_redlined = os.path.join(outdir, "flagged_output_redlined.docx")
    out_docx_highlighted = os.path.join(outdir, "flagged_output_highlighted.docx")
    out_docx_main = os.path.join(outdir, "flagged_output.docx")
    
    # Version 1: Python-Redlines tracked changes (preferred)
    redlined_path = process_docx_with_redlines(input_docx, hits, outdir)
    if redlined_path:
        out_docx_redlined = redlined_path
        # Use redlined version as main output
        import shutil
        shutil.copy2(out_docx_redlined, out_docx_main)
    
    # Version 2: Simple highlighting (fallback)
    highlight_docx_simple(input_docx, hits, out_docx_highlighted)
    
    # If redlines failed, use highlighted version as main
    if not redlined_path:
        import shutil
        shutil.copy2(out_docx_highlighted, out_docx_main)
    
    print(f"Created DOCX versions:")
    if redlined_path:
        print(f"  - Tracked changes (Python-Redlines): {out_docx_redlined}")
    print(f"  - Highlighted: {out_docx_highlighted}")
    print(f"  - Main output: {out_docx_main}")
    
    return out_docx_main, hits

# -----------------------------
# Pipeline
# -----------------------------

def process_file(input_file: str,
                flagged_terms: List[str],
                repl_map: Dict[str, str],
                outdir: str,
                style: str,
                model: str,
                temperature: float,
                api_type: str = "auto",
                skip_terms: List[str] = None) -> Tuple[str, List[Hit]]:
    """Process either PDF or DOCX file based on file extension."""
    file_type = detect_file_type(input_file)
    
    # Extract text for document analysis
    analysis_text = ""
    if file_type == 'pdf':
        doc = fitz.open(input_file)
        for page in doc:
            analysis_text += page.get_text()
        doc.close()
    elif file_type == 'docx':
        doc = Document(input_file)
        for paragraph in doc.paragraphs:
            analysis_text += paragraph.text + "\n"
    
    # Perform document analysis
    print("🔍 Analyzing document themes and strategic alignment...")
    document_analysis = analyze_document_themes(
        pdf_text=analysis_text if file_type == 'pdf' else None,
        docx_text=analysis_text if file_type == 'docx' else None,
        model=model,
        api_type=api_type
    )
    
    # Save document analysis
    analysis_file = os.path.join(outdir, "document_analysis.json")
    with open(analysis_file, 'w', encoding='utf-8') as f:
        json.dump({
            "document_type": document_analysis.document_type,
            "main_themes": document_analysis.main_themes,
            "scientific_focus": document_analysis.scientific_focus,
            "target_population": document_analysis.target_population,
            "intervention_approach": document_analysis.intervention_approach,
            "key_objectives": document_analysis.key_objectives,
            "methodology": document_analysis.methodology,
            "policy_relevance": document_analysis.policy_relevance,
            "nih_priority_alignment": document_analysis.nih_priority_alignment,
            "alignment_opportunities": document_analysis.alignment_opportunities,
            "strategic_recommendations": document_analysis.strategic_recommendations,
            "nih_alignment_score": document_analysis.nih_alignment_score,
            "project_2025_score": document_analysis.project_2025_score,
            "alabama_sb129_score": document_analysis.alabama_sb129_score
        }, f, indent=2, ensure_ascii=False)
    
    print(f"📊 Document Analysis Complete:")
    print(f"   Type: {document_analysis.document_type}")
    print(f"   NIH Alignment Score: {document_analysis.nih_alignment_score:.1f}%")
    print(f"   Project 2025 Score: {document_analysis.project_2025_score:.1f}%")
    print(f"   Alabama SB 129 Score: {document_analysis.alabama_sb129_score:.1f}%")
    print(f"   Key Themes: {', '.join(document_analysis.main_themes[:3])}")
    if document_analysis.nih_priority_alignment:
        print(f"   NIH Priorities: {', '.join(document_analysis.nih_priority_alignment[:3])}")
    
    # Continue with existing processing
    if file_type == 'pdf':
        return process_pdf(input_file, flagged_terms, repl_map, outdir, style, model, temperature, api_type, skip_terms)
    elif file_type == 'docx':
        return process_docx_file(input_file, flagged_terms, repl_map, outdir, model, temperature, api_type, skip_terms)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def process_pdf(input_pdf: str,
                flagged_terms: List[str],
                repl_map: Dict[str, str],
                outdir: str,
                style: str,
                model: str,
                temperature: float,
                api_type: str = "auto",
                skip_terms: List[str] = None) -> Tuple[str, List[Hit]]:
    os.makedirs(outdir, exist_ok=True)
    out_pdf = os.path.join(outdir, "flagged_output.pdf")

    # Build search index
    all_terms = build_match_list(flagged_terms, repl_map)
    phrase_tokens = build_phrase_tokens(all_terms)

    # Determine API type and print info
    if api_type == "gemini" or (api_type == "auto" and _GEMINI_AVAILABLE and os.environ.get("GEMINI_API_KEY")):
        print("Using Gemini 1.5 Flash for language processing")
    else:
        print(f"Using OpenAI {model} for language processing")

    all_hits: List[Hit] = []
    
    # Process terms in parallel chunks
    chunk_size = 20  # Smaller chunks for parallel processing
    total_terms = len(phrase_tokens)
    num_processes = min(4, mp.cpu_count())  # Use up to 4 processes
    
    print(f"Processing {total_terms} terms in chunks of {chunk_size} using {num_processes} processes...")
    
    with fitz.open(input_pdf) as doc:
        total_pages = len(doc)
        print(f"Processing {total_pages} pages with Overkill preset...")
        
        # Extract all page text and words first
        page_texts = []
        page_words_list = []
        for page_num, page in enumerate(doc):
            page_texts.append(page.get_text())
            page_words_list.append(words_by_order(page))
        
        # Process each page with parallel term chunks
        start_time = time.time()
        for page_num, page in enumerate(doc):
            if page_num % 5 == 0:  # Print every 5 pages for more frequent updates
                elapsed = time.time() - start_time
                rate = (page_num + 1) / elapsed if elapsed > 0 else 0
                eta = (total_pages - page_num - 1) / rate if rate > 0 else 0
                print(f"Processing page {page_num + 1}/{total_pages} ({(page_num + 1)/total_pages*100:.1f}%) - Rate: {rate:.1f} pages/sec - ETA: {eta/60:.1f} min")
            
            # Prepare arguments for parallel processing
            page_text = page_texts[page_num]
            page_words = page_words_list[page_num]
            chunk_args = []
            
            for chunk_start in range(0, total_terms, chunk_size):
                chunk_end = min(chunk_start + chunk_size, total_terms)
                chunk_phrase_tokens = phrase_tokens[chunk_start:chunk_end]
                chunk_args.append((page_text, chunk_phrase_tokens, repl_map, model, temperature, page_num + 1, page_words, api_type, skip_terms or []))
            
            # Process chunks in parallel
            with mp.Pool(processes=num_processes) as pool:
                chunk_results = pool.map(process_terms_chunk, chunk_args)
            
            # Collect all hits from this page
            page_hits = []
            for chunk_hits in chunk_results:
                page_hits.extend(chunk_hits)
            
            # Collect hits (annotation will be done later)
            all_hits.extend(page_hits)
            
            if page_hits:
                print(f"  Page {page_num + 1} complete: {len(page_hits)} total matches found")
        
        # Now annotate all hits found
        print(f"Annotating {len(all_hits)} hits found...")
        for hit in all_hits:
            page = doc[hit.page_num - 1]  # Convert to 0-based index
            # Use the proper bounding box calculated during matching
            annotate_hit(page, hit, style=style)
        
        doc.save(out_pdf, deflate=True, garbage=4)

    return out_pdf, all_hits

def export_reports(hits: List[Hit], outdir: str) -> Tuple[str, str]:
    rows = [asdict(h) for h in hits]
    df = pd.DataFrame(rows, columns=["page_num", "original_key", "matched_text", "suggestion", "reason", "bbox", "context"])
    csv_path = os.path.join(outdir, "flag_report.csv")
    json_path = os.path.join(outdir, "flag_report.json")
    df.to_csv(csv_path, index=False)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(rows, f, indent=2, ensure_ascii=False)
    return csv_path, json_path

def process_batch(input_files: List[str], 
                  flagged_terms: List[str], 
                  repl_map: Dict[str, str], 
                  outdir: str, 
                  style: str, 
                  model: str, 
                  temperature: float,
                  api_type: str = "auto",
                  skip_terms: List[str] = None) -> Dict[str, Tuple[str, List[Hit]]]:
    """Process multiple files in batch."""
    results = {}
    
    for input_file in tqdm(input_files, desc="Processing files", unit="file"):
        try:
            print(f"\nProcessing: {input_file}")
            out_file, hits = process_file(
                input_file=input_file,
                flagged_terms=flagged_terms,
                repl_map=repl_map,
                outdir=outdir,
                style=style,
                model=model,
                temperature=temperature,
                api_type=api_type,
                skip_terms=skip_terms
            )
            results[input_file] = (out_file, hits)
            print(f"✅ Completed: {input_file} ({len(hits)} flags found)")
        except Exception as e:
            print(f"❌ Error processing {input_file}: {e}")
            results[input_file] = (None, [])
    
    return results

def main():
    parser = argparse.ArgumentParser(description="Flag, annotate, and LLM-suggest revisions for PDFs and DOCX files.")
    parser.add_argument("input_file", nargs="+", help="Input file(s) (PDF or DOCX) - can specify multiple files for batch processing")
    parser.add_argument("--flagged", required=True, help="Path to flagged_terms.json (list of phrases to flag)")
    parser.add_argument("--map", required=True, help="Path to replacements.json (term -> suggestion)")
    parser.add_argument("--outdir", default="out", help="Output directory")
    parser.add_argument("--style", choices=["highlight", "underline"], default="highlight", help="Annotation style")
    parser.add_argument("--model", default="gpt-4.1-mini", help="OpenAI model to use (ignored if using Gemini)")
    parser.add_argument("--temperature", type=float, default=0.2, help="LLM temperature")
    parser.add_argument("--api", choices=["openai", "gemini", "auto"], default="auto", help="API to use: openai, gemini, or auto (default: auto)")
    parser.add_argument("--skip-terms", nargs="*", default=[], help="Terms to skip during processing (will detect variations)")
    parser.add_argument("--env-file", default=None, help="Path to a .env file containing API keys (optional)")
    args = parser.parse_args()

    # Load environment variables (OPENAI_API_KEY) from .env if present
    load_env(args.env_file)

    flagged_terms = load_json_file(args.flagged)
    repl_map = load_json_file(args.map)

    if not isinstance(flagged_terms, list):
        raise SystemExit("flagged_terms.json must be a JSON list.")
    if not isinstance(repl_map, dict):
        raise SystemExit("replacements.json must be a JSON object mapping 'term' -> 'suggestion'.")

    # Check if batch processing (multiple files)
    if len(args.input_file) > 1:
        print(f"Batch processing {len(args.input_file)} files...")
        results = process_batch(
            input_files=args.input_file,
            flagged_terms=flagged_terms,
            repl_map=repl_map,
            outdir=args.outdir,
            style=args.style,
            model=args.model,
            temperature=args.temperature,
            api_type=args.api,
            skip_terms=args.skip_terms
        )
        
        # Summary
        total_flags = sum(len(hits) for _, (_, hits) in results.items())
        successful = sum(1 for _, (out_file, _) in results.items() if out_file is not None)
        print(f"\n📊 Batch Processing Summary:")
        print(f"  Files processed: {successful}/{len(args.input_file)}")
        print(f"  Total flags found: {total_flags}")
        
    else:
        # Single file processing
        out_file, hits = process_file(
            input_file=args.input_file[0],
            flagged_terms=flagged_terms,
            repl_map=repl_map,
            outdir=args.outdir,
            style=args.style,
            model=args.model,
            temperature=args.temperature,
            api_type=args.api,
            skip_terms=args.skip_terms
        )

        # Export reports
        csv_path, json_path = export_reports(hits, args.outdir)

        print(f"Annotated file: {out_file}")
        print(f"CSV report:     {csv_path}")
        print(f"JSON report:    {json_path}")
        print(f"Total flags:    {len(hits)}")

if __name__ == "__main__":
    main()