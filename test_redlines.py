#!/usr/bin/env python3
"""
Test Python-Redlines functionality
"""

from python_redlines.engines import XmlPowerToolsEngine
from docx import Document

def test_redlines():
    """Test basic Python-Redlines functionality."""
    try:
        # Create a simple test document
        doc = Document()
        doc.add_paragraph("This is a test document with disparities and burden.")
        doc.save("test_original.docx")
        
        # Create a modified version
        doc2 = Document()
        doc2.add_paragraph("This is a test document with differences and challenges.")
        doc2.save("test_modified.docx")
        
        # Use Python-Redlines
        engine = XmlPowerToolsEngine()
        redline_bytes, _, _ = engine.run_redline(
            author_tag="Test Author",
            original="test_original.docx",
            modified="test_modified.docx"
        )
        
        # Save the redlined document
        with open("test_redlined.docx", 'wb') as f:
            f.write(redline_bytes)
        
        print("✅ Python-Redlines test successful!")
        print("Created test_redlined.docx with tracked changes")
        
    except Exception as e:
        print(f"❌ Python-Redlines test failed: {e}")

if __name__ == "__main__":
    test_redlines()
